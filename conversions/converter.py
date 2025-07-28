"""
File conversion services and utilities.
"""
import os
import tempfile
import uuid
from typing import Dict, Any, Optional, Tuple
from io import BytesIO
from pathlib import Path
import mimetypes
from datetime import datetime

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.conf import settings
from django.utils import timezone

# Image processing
try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

# Document processing
try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

# Document libraries
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Office document processing
try:
    import openpyxl
    import xlsxwriter
    import csv
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class ConversionError(Exception):
    """Custom exception for conversion errors."""
    pass


class FileConverterService:
    """Service for handling file conversions."""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def create_conversion_job(self, input_file, output_format: str, user=None, webhook_url=None):
        """
        Create a conversion job in the database.
        
        Args:
            input_file: Uploaded file object
            output_format: Target format for conversion
            user: User object (optional for anonymous conversions)
            webhook_url: URL to notify when conversion is complete (optional)
            
        Returns:
            ConversionJob instance
        """
        from .models import ConversionJob, FileFormat
        
        # Detect input format from file extension
        filename = input_file.name
        input_format = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
        
        # Get format objects
        try:
            input_format_obj = FileFormat.objects.get(name=input_format)
        except FileFormat.DoesNotExist:
            # Create a basic format entry if it doesn't exist
            input_format_obj = FileFormat.objects.create(
                name=input_format,
                category='unknown',
                mime_type=mimetypes.guess_type(filename)[0] or 'application/octet-stream',
                is_input_supported=True,
                is_output_supported=False
            )
        
        try:
            output_format_obj = FileFormat.objects.get(name=output_format)
        except FileFormat.DoesNotExist:
            # Create a basic format entry if it doesn't exist
            output_format_obj = FileFormat.objects.create(
                name=output_format,
                category='unknown',
                mime_type='application/octet-stream',
                is_input_supported=False,
                is_output_supported=True
            )
        
        # Create conversion job
        job_data = {
            'user': user,
            'input_format': input_format_obj,
            'output_format': output_format_obj,
            'input_file': input_file,
            'status': 'pending',
        }
        
        # Only add webhook_url if provided
        if webhook_url:
            job_data['webhook_url'] = webhook_url
            
        job = ConversionJob.objects.create(**job_data)
        return job
    
    def convert_file(self, job_id: str) -> Dict[str, Any]:
        """
        Convert a file using an existing job ID.
        
        Args:
            job_id: UUID of the conversion job
            
        Returns:
            Dictionary with conversion result
        """
        from .models import ConversionJob
        
        try:
            job = ConversionJob.objects.get(id=job_id)
        except ConversionJob.DoesNotExist:
            return {'success': False, 'error': 'Job not found'}
        
        try:
            # Update job status
            job.status = 'processing'
            job.save()
            
            # Get file content and formats
            input_file = job.input_file
            input_format = job.input_format.name
            output_format = job.output_format.name
            
            # Perform conversion
            converted_content, output_filename = self._convert_file_content(
                input_file, input_format, output_format
            )
            
            # Save output file
            output_file = ContentFile(converted_content, name=output_filename)
            job.output_file = output_file
            job.status = 'completed'
            job.completed_at = timezone.now()
            job.save()
            
            return {
                'success': True,
                'job_id': str(job.id),
                'output_filename': output_filename,
                'download_url': f'/api/v1/conversions/jobs/{job.id}/download/'
            }
            
        except Exception as e:
            job.status = 'failed'
            job.error_message = str(e)
            job.save()
            return {'success': False, 'error': str(e)}
    
    def _convert_file_content(self, input_file, input_format: str, output_format: str, 
                    options: Dict[str, Any] = None) -> Tuple[bytes, str]:
        """
        Convert a file from one format to another.
        
        Args:
            input_file: File object or path
            input_format: Input file format (e.g., 'jpg', 'png')
            output_format: Output file format (e.g., 'png', 'pdf')
            options: Conversion options
            
        Returns:
            Tuple of (converted_file_bytes, output_filename)
            
        Raises:
            ConversionError: If conversion fails
        """
        if options is None:
            options = {}
            
        # Normalize formats
        input_format = input_format.lower().strip('.')
        output_format = output_format.lower().strip('.')
        
        # Get file content
        if hasattr(input_file, 'read'):
            file_content = input_file.read()
            if hasattr(input_file, 'name'):
                filename = input_file.name
            else:
                filename = f"input.{input_format}"
        else:
            with open(input_file, 'rb') as f:
                file_content = f.read()
            filename = os.path.basename(input_file)
        
        # Generate output filename
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}.{output_format}"
        
        try:
            # Route to appropriate converter
            if self._is_image_conversion(input_format, output_format):
                return self._convert_image(file_content, input_format, output_format, options), output_filename
            elif self._is_document_conversion(input_format, output_format):
                return self._convert_document(file_content, input_format, output_format, options), output_filename
            elif self._is_spreadsheet_conversion(input_format, output_format):
                return self._convert_spreadsheet(file_content, input_format, output_format, options), output_filename
            elif self._is_text_conversion(input_format, output_format):
                return self._convert_text(file_content, input_format, output_format, options), output_filename
            else:
                raise ConversionError(f"Conversion from {input_format} to {output_format} is not supported yet")
                
        except Exception as e:
            raise ConversionError(f"Conversion failed: {str(e)}")
    
    def _is_image_conversion(self, input_format: str, output_format: str) -> bool:
        """Check if this is an image conversion."""
        image_formats = {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp', 'ico'}
        return input_format in image_formats and output_format in image_formats
    
    def _is_document_conversion(self, input_format: str, output_format: str) -> bool:
        """Check if this is a document conversion."""
        doc_formats = {'pdf', 'docx', 'doc', 'txt', 'rtf', 'odt', 'html'}
        return input_format in doc_formats and output_format in doc_formats
    
    def _is_spreadsheet_conversion(self, input_format: str, output_format: str) -> bool:
        """Check if this is a spreadsheet conversion."""
        sheet_formats = {'xlsx', 'xls', 'csv', 'ods', 'tsv'}
        return input_format in sheet_formats and output_format in sheet_formats
    
    def _is_text_conversion(self, input_format: str, output_format: str) -> bool:
        """Check if this is a text/data conversion."""
        text_formats = {'txt', 'json', 'xml', 'yaml', 'yml', 'csv', 'tsv'}
        return input_format in text_formats and output_format in text_formats
    
    def _convert_image(self, file_content: bytes, input_format: str, 
                      output_format: str, options: Dict[str, Any]) -> bytes:
        """Convert image files."""
        if not PILLOW_AVAILABLE:
            raise ConversionError("Pillow library is required for image conversions")
        
        try:
            # Open image from bytes
            image = Image.open(BytesIO(file_content))
            
            # Handle transparency for formats that don't support it
            if output_format.lower() in ['jpg', 'jpeg'] and image.mode in ['RGBA', 'LA']:
                # Create a white background
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'RGBA':
                    background.paste(image, mask=image.split()[-1])  # Use alpha channel as mask
                else:
                    background.paste(image)
                image = background
            
            # Apply options
            quality = options.get('quality', 95)
            if 'width' in options or 'height' in options:
                width = options.get('width')
                height = options.get('height')
                if width and height:
                    image = image.resize((width, height), Image.Resampling.LANCZOS)
                elif width:
                    ratio = width / image.width
                    height = int(image.height * ratio)
                    image = image.resize((width, height), Image.Resampling.LANCZOS)
                elif height:
                    ratio = height / image.height
                    width = int(image.width * ratio)
                    image = image.resize((width, height), Image.Resampling.LANCZOS)
            
            # Save to bytes
            output_buffer = BytesIO()
            save_format = output_format.upper()
            if save_format == 'JPG':
                save_format = 'JPEG'
            
            if save_format in ['JPEG', 'WEBP']:
                image.save(output_buffer, format=save_format, quality=quality, optimize=True)
            else:
                image.save(output_buffer, format=save_format)
            
            return output_buffer.getvalue()
            
        except Exception as e:
            raise ConversionError(f"Image conversion failed: {str(e)}")
    
    def _convert_document(self, file_content: bytes, input_format: str,
                         output_format: str, options: Dict[str, Any]) -> bytes:
        """Convert document files using proper libraries."""
        
        # Text to DOCX conversion
        if input_format == 'txt' and output_format == 'docx':
            return self._txt_to_docx(file_content)
        
        # DOCX to Text conversion  
        elif input_format == 'docx' and output_format == 'txt':
            return self._docx_to_txt(file_content)
        
        # Text to HTML conversion
        elif input_format == 'txt' and output_format == 'html':
            return self._txt_to_html(file_content)
        
        # HTML to Text conversion
        elif input_format == 'html' and output_format == 'txt':
            return self._html_to_txt(file_content)
        
        # Text to PDF conversion
        elif input_format == 'txt' and output_format == 'pdf':
            return self._txt_to_pdf(file_content)
        
        # PDF to Text conversion
        elif input_format == 'pdf' and output_format == 'txt':
            return self._pdf_to_txt(file_content)
        
        # Same format (just copy)
        elif input_format == output_format:
            return file_content
        
        else:
            raise ConversionError(f"Document conversion from {input_format} to {output_format} is not yet implemented")
    
    def _txt_to_docx(self, file_content: bytes) -> bytes:
        """Convert text file to DOCX using python-docx."""
        try:
            from docx import Document
            from io import BytesIO
            
            # Decode text content
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Converted Document', 0)
            
            # Add content - split by lines and paragraphs
            paragraphs = text_content.split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Add paragraph
                    p = doc.add_paragraph()
                    lines = paragraph_text.split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            if i > 0:
                                p.add_run('\n')
                            p.add_run(line.strip())
            
            # If no content was added, add a default message
            if len(doc.paragraphs) <= 1:  # Only title paragraph
                doc.add_paragraph("This document was converted from a text file.")
                doc.add_paragraph(text_content[:1000] + "..." if len(text_content) > 1000 else text_content)
            
            # Save to bytes
            output_buffer = BytesIO()
            doc.save(output_buffer)
            return output_buffer.getvalue()
            
        except Exception as e:
            raise ConversionError(f"Text to DOCX conversion failed: {str(e)}")
    
    def _docx_to_txt(self, file_content: bytes) -> bytes:
        """Convert DOCX file to text using python-docx."""
        try:
            from docx import Document
            from io import BytesIO
            
            # Load document from bytes
            doc = Document(BytesIO(file_content))
            
            # Extract text
            text_lines = []
            for paragraph in doc.paragraphs:
                text_lines.append(paragraph.text)
            
            # Join with newlines
            text_content = '\n'.join(text_lines)
            
            # If no content, provide a message
            if not text_content.strip():
                text_content = "No readable text content found in the document."
            
            return text_content.encode('utf-8')
            
        except Exception as e:
            raise ConversionError(f"DOCX to text conversion failed: {str(e)}")
    
    def _txt_to_html(self, file_content: bytes) -> bytes:
        """Convert text to HTML."""
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Escape HTML characters
            import html
            escaped_text = html.escape(text_content)
            
            # Convert line breaks to HTML
            html_content = escaped_text.replace('\n\n', '</p><p>').replace('\n', '<br>')
            
            # Wrap in HTML structure
            full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; line-height: 1.6; }}
        p {{ margin-bottom: 1em; }}
    </style>
</head>
<body>
    <h1>Converted Document</h1>
    <p>{html_content}</p>
</body>
</html>"""
            return full_html.encode('utf-8')
            
        except Exception as e:
            raise ConversionError(f"Text to HTML conversion failed: {str(e)}")
    
    def _html_to_txt(self, file_content: bytes) -> bytes:
        """Convert HTML to plain text."""
        try:
            import re
            import html
            
            html_content = file_content.decode('utf-8', errors='ignore')
            
            # Remove script and style elements
            html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Replace common HTML elements with text equivalents
            html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'<p\s*[^>]*>', '\n\n', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'</p>', '', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'<h[1-6]\s*[^>]*>', '\n\n', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'</h[1-6]>', '\n', html_content, flags=re.IGNORECASE)
            
            # Remove all remaining HTML tags
            text_content = re.sub(r'<[^>]+>', '', html_content)
            
            # Decode HTML entities
            text_content = html.unescape(text_content)
            
            # Clean up whitespace
            text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
            text_content = text_content.strip()
            
            return text_content.encode('utf-8')
            
        except Exception as e:
            raise ConversionError(f"HTML to text conversion failed: {str(e)}")
    
    def _txt_to_pdf(self, file_content: bytes) -> bytes:
        """Convert text to PDF using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from io import BytesIO
            
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            normal_style = styles['Normal']
            
            # Create story (content)
            story = []
            
            # Add title
            story.append(Paragraph("Converted Document", title_style))
            story.append(Spacer(1, 12))
            
            # Add content paragraphs
            paragraphs = text_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Clean the text for PDF
                    clean_text = para_text.replace('\n', '<br/>').strip()
                    story.append(Paragraph(clean_text, normal_style))
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            return buffer.getvalue()
            
        except Exception as e:
            raise ConversionError(f"Text to PDF conversion failed: {str(e)}")
    
    def _pdf_to_txt(self, file_content: bytes) -> bytes:
        """Convert PDF to text with enhanced structure detection."""
        try:
            # Try enhanced structure-preserving extraction first
            from .pdf_structure_extractor import extract_structured_pdf_text
            
            try:
                structured_text, document_structure = extract_structured_pdf_text(
                    file_content, 
                    options={'preserve_structure': True}
                )
                
                # Add document structure metadata as comments
                metadata = [
                    f"# Document Structure Analysis",
                    f"# Total Pages: {document_structure.get('total_pages', 'Unknown')}",
                    f"# Headers: {document_structure.get('statistics', {}).get('headers', 0)}",
                    f"# Paragraphs: {document_structure.get('statistics', {}).get('paragraphs', 0)}",
                    f"# Lists: {document_structure.get('statistics', {}).get('lists', 0)}",
                    f"# Tables: {document_structure.get('statistics', {}).get('tables', 0)}",
                    f"# Extraction Method: Enhanced Structure Detection",
                    f"# {'='*60}",
                    ""
                ]
                
                final_text = "\n".join(metadata) + "\n" + structured_text
                return final_text.encode('utf-8')
                
            except Exception as enhanced_error:
                print(f"Enhanced extraction failed: {enhanced_error}, falling back to basic extraction")
                # Fall back to basic PyPDF2 extraction
                pass
            
            # Fallback to basic PyPDF2 extraction
            import PyPDF2
            from io import BytesIO
            
            # Read PDF from bytes
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            
            # Extract text from all pages with basic structure
            text_content = ""
            text_content += "# Document Extracted with Basic Method\n"
            text_content += f"# Total Pages: {len(pdf_reader.pages)}\n"
            text_content += "# Note: For better structure detection, install pdfplumber and pymupdf\n"
            text_content += f"# {'='*60}\n\n"
            
            for page_num in range(len(pdf_reader.pages)):
                text_content += f"\n{'='*50}\n"
                text_content += f"PAGE {page_num + 1}\n"
                text_content += f"{'='*50}\n\n"
                
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text:
                    # Apply basic structure formatting
                    lines = page_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Basic header detection
                        if (len(line) < 100 and 
                            (line.isupper() or 
                             any(line.startswith(prefix) for prefix in ['1.', '2.', '3.', '4.', '5.']) or
                             any(line.startswith(prefix) for prefix in ['I.', 'II.', 'III.', 'IV.', 'V.']))):
                            text_content += f"\n# {line}\n\n"
                        
                        # Basic list detection
                        elif (line.startswith(('• ', '* ', '- ', '+ ')) or
                              any(line.startswith(f"{i}.") for i in range(1, 20)) or
                              any(line.startswith(f"{chr(97+i)}.") for i in range(26))):
                            text_content += f"  • {line.lstrip('•*-+ ').lstrip('0123456789.).abcdefghijklmnopqrstuvwxyz')}\n"
                        
                        else:
                            text_content += f"{line}\n"
                    
                    text_content += "\n"
            
            # Clean up the text
            text_content = text_content.strip()
            
            if not text_content or text_content.count('#') == text_content.count('\n'):
                text_content = "No readable text found in the PDF document."
            
            return text_content.encode('utf-8')
            
        except Exception as e:
            raise ConversionError(f"PDF to text conversion failed: {str(e)}")
    
    def _convert_spreadsheet(self, file_content: bytes, input_format: str,
                           output_format: str, options: Dict[str, Any]) -> bytes:
        """Convert spreadsheet files."""
        if not EXCEL_AVAILABLE:
            raise ConversionError("Excel processing libraries are required for spreadsheet conversions")
        
        try:
            if input_format == 'csv' and output_format == 'xlsx':
                # CSV to Excel
                import io
                import pandas as pd
                
                # Read CSV
                csv_data = file_content.decode('utf-8', errors='ignore')
                df = pd.read_csv(io.StringIO(csv_data))
                
                # Write to Excel
                output_buffer = BytesIO()
                with pd.ExcelWriter(output_buffer, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Sheet1')
                
                return output_buffer.getvalue()
            
            elif input_format == 'xlsx' and output_format == 'csv':
                # Excel to CSV
                import pandas as pd
                
                # Read Excel
                df = pd.read_excel(BytesIO(file_content))
                
                # Write to CSV
                output_buffer = io.StringIO()
                df.to_csv(output_buffer, index=False)
                
                return output_buffer.getvalue().encode('utf-8')
            
            else:
                raise ConversionError(f"Spreadsheet conversion from {input_format} to {output_format} not implemented yet")
                
        except Exception as e:
            raise ConversionError(f"Spreadsheet conversion failed: {str(e)}")
    
    def _convert_text(self, file_content: bytes, input_format: str,
                     output_format: str, options: Dict[str, Any]) -> bytes:
        """Convert text/data files."""
        try:
            if input_format == 'json' and output_format == 'yaml':
                import json
                import yaml
                
                # Parse JSON
                data = json.loads(file_content.decode('utf-8'))
                
                # Convert to YAML
                yaml_content = yaml.dump(data, default_flow_style=False)
                return yaml_content.encode('utf-8')
            
            elif input_format in ['yaml', 'yml'] and output_format == 'json':
                import json
                import yaml
                
                # Parse YAML
                data = yaml.safe_load(file_content.decode('utf-8'))
                
                # Convert to JSON
                json_content = json.dumps(data, indent=2)
                return json_content.encode('utf-8')
            
            elif input_format == 'csv' and output_format == 'json':
                import csv
                import json
                import io
                
                # Parse CSV
                csv_data = file_content.decode('utf-8', errors='ignore')
                reader = csv.DictReader(io.StringIO(csv_data))
                data = list(reader)
                
                # Convert to JSON
                json_content = json.dumps(data, indent=2)
                return json_content.encode('utf-8')
            
            else:
                raise ConversionError(f"Text conversion from {input_format} to {output_format} not implemented yet")
                
        except Exception as e:
            raise ConversionError(f"Text conversion failed: {str(e)}")
    
    def get_file_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Get information about a file."""
        info = {
            'size': len(file_content),
            'filename': filename,
            'extension': os.path.splitext(filename)[1].lower().strip('.'),
        }
        
        # Try to determine MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            info['mime_type'] = mime_type
        
        # For images, get dimensions
        if PILLOW_AVAILABLE and info['extension'] in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']:
            try:
                image = Image.open(BytesIO(file_content))
                info['width'] = image.width
                info['height'] = image.height
                info['mode'] = image.mode
            except:
                pass
        
        return info


# Global converter instance
converter_service = FileConverterService()
